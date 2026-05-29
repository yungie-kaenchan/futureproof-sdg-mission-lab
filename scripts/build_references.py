#!/usr/bin/env python3
"""Build a standalone one-page, two-column APA-7 reference list (.docx → PDF)."""
import os, sys
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
import fp_docgen as fp
from docx.shared import Pt, Mm
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = "docs/judge-pack"; COVERS = "docs/document-design/covers"

REFS = [
 "Anderson, L. W., & Krathwohl, D. R. (Eds.). (2001). A taxonomy for learning, teaching, and assessing: A revision of Bloom’s taxonomy of educational objectives. New York: Longman.",
 "Ausubel, D. P. (1968). Educational psychology: A cognitive view. New York: Holt, Rinehart and Winston.",
 "Borg, W. R., & Gall, M. D. (1989). Educational research: An introduction (5th ed.). New York: Longman.",
 "Bruner, J. S. (1960). The process of education. Cambridge, MA: Harvard University Press.",
 "Bruner, J. S. (1966). Toward a theory of instruction. Cambridge, MA: Harvard University Press.",
 "Bruner, J. S. (1996). The culture of education. Cambridge, MA: Harvard University Press.",
 "CAST. (2024). Universal Design for Learning Guidelines (Version 3.0). Lynnfield, MA: CAST. https://udlguidelines.cast.org",
 "Churches, A. (2008). Bloom’s Digital Taxonomy. Educational Origami. https://www.eduteka.icesi.edu.co/pdfdir/TaxonomiaBloomDigital.pdf",
 "Council of Europe. (2020). Common European Framework of Reference for Languages: Learning, teaching, assessment — Companion volume. Strasbourg: Council of Europe Publishing.",
 "Coyle, D., Hood, P., & Marsh, D. (2010). CLIL: Content and language integrated learning. Cambridge, UK: Cambridge University Press.",
 "Flavell, J. H. (1979). Metacognition and cognitive monitoring: A new area of cognitive–developmental inquiry. American Psychologist, 34(10), 906–911. https://doi.org/10.1037/0003-066X.34.10.906",
 "Hattie, J. (2009). Visible learning: A synthesis of over 800 meta-analyses relating to achievement. London: Routledge.",
 "Hattie, J. (2012). Visible learning for teachers: Maximizing impact on learning. London: Routledge.",
 "Kimmons, R., Graham, C. R., & West, R. E. (2020). The PICRAT model for technology integration in teacher preparation. Contemporary Issues in Technology and Teacher Education, 20(1), 176–198.",
 "Krashen, S. D. (1982). Principles and practice in second language acquisition. Oxford: Pergamon.",
 "Krashen, S. D. (1985). The input hypothesis: Issues and implications. London: Longman.",
 "Lave, J., & Wenger, E. (1991). Situated learning: Legitimate peripheral participation. Cambridge, UK: Cambridge University Press.",
 "Mayer, R. E. (2009). Multimedia learning (2nd ed.). Cambridge, UK: Cambridge University Press.",
 "Mayer, R. E. (Ed.). (2014). The Cambridge handbook of multimedia learning (2nd ed.). Cambridge, UK: Cambridge University Press.",
 "Mishra, P., & Koehler, M. J. (2006). Technological pedagogical content knowledge: A framework for teacher knowledge. Teachers College Record, 108(6), 1017–1054.",
 "Mishra, P., Warr, M., & Islam, R. (2023). TPACK in the age of ChatGPT and generative AI. Journal of Digital Learning in Teacher Education, 39(4), 235–251. https://doi.org/10.1080/21532974.2023.2247480",
 "Partnership for 21st Century Learning. (2019). Framework for 21st century learning: Definitions. Hilliard, OH: Battelle for Kids. https://www.battelleforkids.org/networks/p21",
 "Paul, R., & Elder, L. (2007). The art of Socratic questioning: A companion to: The thinker’s guide to analytic thinking. Tomales, CA: Foundation for Critical Thinking.",
 "Pavlov, I. P. (1927). Conditioned reflexes: An investigation of the physiological activity of the cerebral cortex (G. V. Anrep, Trans.). Oxford, UK: Oxford University Press.",
 "Piaget, J. (1970). Genetic epistemology (E. Duckworth, Trans.). New York: Columbia University Press.",
 "Skinner, B. F. (1953). Science and human behavior. New York: Free Press.",
 "Skinner, B. F. (1957). Verbal behavior. New York: Appleton-Century-Crofts.",
 "Sweller, J. (1988). Cognitive load during problem solving: Effects on learning. Cognitive Science, 12(2), 257–285.",
 "Sweller, J., van Merriënboer, J. J. G., & Paas, F. (2019). Cognitive architecture and instructional design: 20 years later. Educational Psychology Review, 31, 261–292. https://doi.org/10.1007/s10648-019-09465-5",
 "Times Higher Education. (2024). Impact Rankings 2024: Methodology. London: THE. https://www.timeshighereducation.com/impactrankings",
 "UNESCO. (2020). Education for sustainable development: A roadmap (ESD for 2030). Paris: UNESCO. https://unesdoc.unesco.org/ark:/48223/pf0000374802",
 "Vygotsky, L. S. (1978). Mind in society: The development of higher psychological processes (M. Cole, V. John-Steiner, S. Scribner, & E. Souberman, Eds.). Cambridge, MA: Harvard University Press.",
 "W3C. (2018). Web Content Accessibility Guidelines (WCAG) 2.1. World Wide Web Consortium. https://www.w3.org/TR/WCAG21/",
 "Watson, J. B. (1913). Psychology as the behaviorist views it. Psychological Review, 20, 158–177.",
 "Wenger, E. (1998). Communities of practice: Learning, meaning, and identity. Cambridge, UK: Cambridge University Press.",
 "พระราชบัญญัติคุ้มครองข้อมูลส่วนบุคคล พ.ศ. ๒๕๖๒. (๒๕๖๒, ๒๗ พฤษภาคม). ราชกิจจานุเบกษา, ๑๓๖(๖๙ ก), ๕๒–๙๖.",
]

def set_two_columns(section, num=2, space_twips=300):
    cols = section._sectPr.find(qn('w:cols'))
    if cols is None:
        cols = OxmlElement('w:cols'); section._sectPr.append(cols)
    cols.set(qn('w:num'), str(num)); cols.set(qn('w:space'), str(space_twips))

doc = fp.new_doc()
s0 = doc.sections[0]
for a in ('top_margin', 'bottom_margin', 'left_margin', 'right_margin'):
    setattr(s0, a, Mm(15))

# full-width title (1 column)
fp.kicker(doc, "FUTUREPROOF · SDG Mission Journey · Reference List · APA 7")
pt = doc.add_paragraph(); pt.paragraph_format.space_after = Pt(7)
fp._add_run(pt, "วรรณกรรมอ้างอิง / References", font=fp.SERIF, size=16, color=fp.DRED, cs=fp.THAI, ea=fp.THAI)
fp.bottom_rule(pt, fp.GOLDH, sz=6)

# continuous section → two columns for the entries
sec2 = doc.add_section(WD_SECTION.CONTINUOUS)
for a in ('top_margin', 'bottom_margin', 'left_margin', 'right_margin'):
    setattr(sec2, a, Mm(15))
set_two_columns(sec2, 2, 300)

for ref in REFS:
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.left_indent = Mm(4.5); pf.first_line_indent = Mm(-4.5)
    pf.space_after = Pt(4); pf.line_spacing = 1.04
    fp._add_run(p, ref, font=fp.SANS, size=8, color=fp.INK, cs=fp.THAI, ea=fp.THAI)

# light text footer (no badge image → keeps the appendix lean)
f = doc.sections[-1].footer; f.is_linked_to_previous = False
fp_p = f.paragraphs[0]; fp_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
fp._add_run(fp_p, "FUTUREPROOF · Reference List (APA 7) · Dr. Payungsak Kaenchan · Faculty of Liberal Arts, Mahidol University · CC BY-NC",
            font=fp.MONO, size=6.8, color=fp.GRAY, ls=0.3)

docx = os.path.join(OUT, "REFERENCE-LIST-APA7.docx")
doc.save(docx); print("saved", docx, "·", len(REFS), "refs")
print("pdf", fp.to_pdf(docx))
