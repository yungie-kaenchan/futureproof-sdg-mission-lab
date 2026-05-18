#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
FUTUREPROOF — Asset-Production Prompt Master generator (v2).

Emits TWO artifacts from one source of truth (the PROMPTS data below):
  • scenarios/PRODUCTION-PROMPTS-MASTER-v2.docx   (hand to producers)
  • scenarios/PRODUCTION-PROMPTS-MASTER-v2.md     (in-repo, diffable)

Covers the TWO fully-built missions exactly as the code wires them, so a
produced file dropped at the stated path auto-wires with zero code change.

The audio SCRIPTS are the verbatim on-screen transcripts pulled from
src/scenarios/*-content.js — the spoken audio MUST match the wired
captions/transcript word-for-word. Do not paraphrase them.

Regenerate:  python3 scripts/gen-production-prompts.py
"""

import os
from datetime import date

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
OUT_DOCX = os.path.join(ROOT, "scenarios", "PRODUCTION-PROMPTS-MASTER-v2.docx")
OUT_MD = os.path.join(ROOT, "scenarios", "PRODUCTION-PROMPTS-MASTER-v2.md")


def wpm(words, seconds):
    return round(words / (seconds / 60.0))


# ── Verbatim transcripts (from src/scenarios/*-content.js — DO NOT EDIT
#    here without changing the content module + .vtt in lockstep) ──
KK = [
    dict(n=1, role="Smallholder rice farmer", th="เกษตรกรผู้ปลูกข้าวรายย่อย",
         path="/assets/scenarios/sdg06-khonkaen/audio/01-rice-farmer.mp3",
         portrait="/assets/scenarios/sdg06-khonkaen/images/stakeholder-01.svg",
         dur=32, gender="male", age="55–65", flags="vulnerable",
         voice="Weathered, unhurried Northeastern (Isan) Thai man speaking "
               "English with a natural Isan-Thai accent. Dignified, tired but "
               "not defeated; the calm of someone stating a plain fact.",
         tone="Quiet conviction. No anger, no pleading. A short pause after "
              "'Now?' and before 'That is just how the ground works.'",
         script="My family has farmed this land for three generations. The "
                "shallow well in our yard — it used to run all year. Now? By "
                "March, it gives nothing. I hear the province will dig deeper "
                "wells in town. I do not blame them. People need to drink. But "
                "the water under our feet is the same water. If they pull more, "
                "mine runs dry sooner. That is just how the ground works."),
    dict(n=2, role="PWA Operations Manager", th="ผู้จัดการฝ่ายปฏิบัติการ การประปาส่วนภูมิภาคขอนแก่น",
         path="/assets/scenarios/sdg06-khonkaen/audio/02-pwa-manager.mp3",
         portrait="/assets/scenarios/sdg06-khonkaen/images/stakeholder-02.svg",
         dur=32, gender="female", age="40–50", flags="institutional",
         voice="Composed, professional Thai woman, central-Thai English, "
               "administrative register. Measured authority, not cold.",
         tone="Brisk and factual; a firm landing on the final three words "
              "'But taps must run.'",
         script="Our statutory mandate is clear: continuous potable water to "
                "around two hundred thousand residents in this service area. In "
                "each of the last three drought years, we came within days of "
                "supply interruption. The twelve-well expansion is sized to "
                "provide a residual margin during a one-in-fifty-year dry "
                "season. We have heard the concerns from rural districts and "
                "from the university working paper. We are prepared to "
                "commission additional monitoring. But taps must run."),
    dict(n=3, role="Food-processing facility Production Director",
         th="ผู้อำนวยการฝ่ายผลิต โรงงานแปรรูปอาหารในแอ่งขอนแก่น",
         path="/assets/scenarios/sdg06-khonkaen/audio/03-plant-director.mp3",
         portrait="/assets/scenarios/sdg06-khonkaen/images/stakeholder-03.svg",
         dur=30, gender="male", age="45–55", flags="private",
         voice="Confident corporate Thai man, polished central-Thai English, "
               "the cadence of a prepared statement. Reasonable, guarded.",
         tone="Controlled and lawyerly; slight emphasis on 'good faith' and "
              "'predictable regulatory environment'.",
         script="Our facility holds a groundwater extraction permit issued "
                "under the framework that has governed industrial water use in "
                "this region for nearly five decades. We extract from the deep "
                "confined layer, well below any household well. Our extraction "
                "volume is monitored and reported quarterly. We support "
                "sustainable allocation. We do not support retroactive "
                "cancellation of permits granted in good faith. A predictable "
                "regulatory environment is essential."),
    dict(n=4, role="Community Health Volunteer (อสม.)",
         th="อาสาสมัครสาธารณสุขประจำหมู่บ้าน",
         path="/assets/scenarios/sdg06-khonkaen/audio/04-health-volunteer.mp3",
         portrait="/assets/scenarios/sdg06-khonkaen/images/stakeholder-04.svg",
         dur=34, gender="female", age="45–58", flags="vulnerable-adjacent",
         voice="Warm, grounded Isan Thai woman, English with an Isan accent, "
               "the voice of a trusted village figure who has seen a lot.",
         tone="Gentle, sincere, rising to a quiet plea on the last sentence "
              "'Please remember the children when you decide.'",
         script="I have walked the houses in my sub-district for eleven years. "
                "When the shallow wells fail in March, families turn to ponds "
                "or river water. Within two weeks, mothers bring children with "
                "diarrhoea to our health post. Older residents stop drinking "
                "enough because they do not trust the taste. Dehydration in the "
                "elderly is a quieter problem, but it is real. Please remember "
                "the children when you decide."),
]

CM = [
    dict(n=1, role="Upland maize smallholder", th="เกษตรกรปลูกข้าวโพดบนพื้นที่สูง",
         path="/assets/scenarios/sdg13-chiangmai/audio/01-maize-grower.mp3",
         vtt="/assets/scenarios/sdg13-chiangmai/audio/01-maize-grower.vtt",
         portrait="/assets/scenarios/sdg13-chiangmai/images/stakeholder-01.svg",
         dur=33, gender="male", age="35–48", flags="vulnerable",
         voice="Northern Thai (kham mueang–inflected) man, English with a "
               "northern-Thai highland accent. Hardworking, plain-spoken, "
               "carrying quiet frustration that does not tip into anger.",
         tone="Steady and direct. A hard, deliberate landing on the last two "
              "sentences: 'I will just do it at night. … A fine is not a choice.'",
         script="I know the smoke is bad. My own children cough too — we "
                "breathe it first, up here, before it reaches the city. But the "
                "buyer sets the date. The field must be clear in days or I lose "
                "the contract, and that contract is the only money my family "
                "sees all year. The machine to bury the stalks costs more than "
                "I earn in two seasons. If you fine me, I will still clear the "
                "field. I will just do it at night. Give me a real choice and I "
                "will take it. A fine is not a choice."),
    dict(n=2, role="Respiratory clinician, public hospital",
         th="แพทย์ระบบทางเดินหายใจ โรงพยาบาลรัฐ",
         path="/assets/scenarios/sdg13-chiangmai/audio/02-clinician.mp3",
         vtt="/assets/scenarios/sdg13-chiangmai/audio/02-clinician.vtt",
         portrait="/assets/scenarios/sdg13-chiangmai/images/stakeholder-02.svg",
         dur=32, gender="female", age="38–48", flags="institutional",
         voice="Precise, caring Thai woman doctor, clear central-Thai English, "
               "clinical exactness softened by real concern.",
         tone="Controlled urgency; the data stated flatly, then a direct "
              "appeal on 'count the children in my ward as stakeholders too.'",
         script="Every burning season I see the same ward fill up. Children on "
                "nebulisers, elderly patients whose oxygen falls for weeks, "
                "asthma that should be stable and is not. PM2.5 at the levels "
                "we record is not a discomfort — it is a measurable rise in "
                "admissions and, in the most fragile, in deaths. I am not "
                "asking you to ignore the farmers. I am asking you to count the "
                "children in my ward as stakeholders too."),
    dict(n=3, role="Agribusiness maize-procurement manager",
         th="ผู้จัดการฝ่ายจัดซื้อข้าวโพด บริษัทเกษตรอุตสาหกรรม",
         path="/assets/scenarios/sdg13-chiangmai/audio/03-procurement.mp3",
         vtt="/assets/scenarios/sdg13-chiangmai/audio/03-procurement.vtt",
         portrait="/assets/scenarios/sdg13-chiangmai/images/stakeholder-03.svg",
         dur=30, gender="male", age="40–52", flags="private",
         voice="Smooth corporate Thai man, polished central-Thai English, the "
               "even delivery of someone used to deflecting responsibility "
               "politely.",
         tone="Measured, non-defensive on the surface; faint firmness on 'We "
              "move when the incentives move.'",
         script="Our contracts specify volume and delivery windows because the "
                "feed mills downstream run on a schedule we do not control. We "
                "do not tell any grower to burn — that is their field decision. "
                "We are open to a certified no-burn supply line; we have "
                "piloted one. But it needs a price premium the market has not "
                "agreed, and it needs the province to fund the machinery gap, "
                "not the buyer alone. We move when the incentives move."),
    dict(n=4, role="District enforcement officer (haze task force)",
         th="เจ้าหน้าที่บังคับใช้กฎหมายระดับอำเภอ ชุดเฉพาะกิจหมอกควัน",
         path="/assets/scenarios/sdg13-chiangmai/audio/04-officer.mp3",
         vtt="/assets/scenarios/sdg13-chiangmai/audio/04-officer.vtt",
         portrait="/assets/scenarios/sdg13-chiangmai/images/stakeholder-04.svg",
         dur=34, gender="female", age="35–45", flags="institutional",
         voice="Frank, slightly weary Thai woman official, northern-inflected "
               "central-Thai English, the candour of someone telling an "
               "uncomfortable truth to a committee.",
         tone="Honest and plain, a little tired; emphasis on 'the last step "
              "after a real alternative — not the first step instead of one.'",
         script="I am the one who knocks on the door after the satellite flags "
                "a hotspot. Honestly: a blanket ban with fines and nothing "
                "else makes my job harder. People stop talking to us. Fires "
                "move to the night. Last season we wrote penalties we could "
                "not collect, and we lost cooperation we spent years building. "
                "Enforcement works as the last step after a real alternative — "
                "not the first step instead of one."),
]

NEG = ("photoreal face of a real identifiable person, celebrity likeness, "
       "logos, brand marks, text, watermark, distorted hands, extra fingers, "
       "low-res, oversaturated, stereotyped 'poverty' tropes, pity framing")


def audio_block(s, mission_slug):
    words = len(s["script"].split())
    pace = wpm(words, s["dur"])
    vtt_line = ("\n• Caption file (must exist, same words): "
                + s.get("vtt", "(see Khon Kaen — .vtt already shipped: "
                  + s["path"].replace(".mp3", ".vtt") + ")")
                ) if True else ""
    return {
        "title": f'A{s["n"]}. Stakeholder {s["n"]:02d} — {s["role"]}',
        "rows": [
            ("Thai role label", s["th"]),
            ("Tool", "ElevenLabs (Multilingual v2 or v3) — English narration, Thai-accented"),
            ("WIRE TO (exact path — drop file here, auto-wires)", s["path"]),
            ("Caption (.vtt) pairing", s.get("vtt", s["path"].replace(".mp3", ".vtt")
                + "  (already shipped for Khon Kaen — audio must match it verbatim)")),
            ("Portrait pairs with", s["portrait"]),
            ("Target duration", f'{s["dur"]} s  (script ≈ {words} words → ~{pace} wpm)'),
            ("Voice profile", s["voice"]),
            ("Performance direction", s["tone"]),
            ("ElevenLabs settings (start here, tune by ear)",
             "Stability 0.40–0.55 · Similarity 0.75 · Style 0.30–0.45 · "
             "Speaker-boost ON. Read at the wpm above; do not rush the final "
             "sentence. One natural breath at each em-dash / question mark."),
        ],
        "script": s["script"],
        "qa": [
            "Spoken words EXACTLY match the script (and the .vtt) — no "
            "paraphrase, no added 'um', no dropped clause.",
            f'Length within ±3 s of {s["dur"]} s.',
            "Accent reads as authentic Thai-English for the role, never "
            "cartoonish or mocking.",
            "Loudness normalised to ~ -16 LUFS; no clipping; clean tail.",
            f'Exported as MP3 (mono, 128 kbps OK) and placed exactly at '
            f'{s["path"]} — filename is load-bearing, do not rename.',
        ],
    }


def portrait_block(n, role, path, flags, palette):
    return {
        "title": f'B{n}. Portrait {n:02d} — {role}',
        "rows": [
            ("Tool", "ChatGPT Image / Midjourney / SDXL — stylised editorial portrait"),
            ("WIRE TO (exact path)", path),
            ("Note on format", "Code references a .svg path. Either (a) export "
             "the render as SVG-wrapped raster at this exact filename, or (b) "
             "save as .png/.webp and change only the file extension in "
             "src/scenarios/<mission>-content.js portrait field. Filename stem "
             "must stay the same."),
            ("Prompt", f"Stylised, dignified editorial portrait illustration of "
             f"a {role.lower()} — a COMPOSITE role, not a real person. "
             f"Three-quarter view, calm direct gaze, neutral studio background, "
             f"{palette} palette consistent with a serious editorial console "
             f"UI. Mature, respectful, agency and competence in the face — "
             f"never pitiable. Subtle Thai contextual cues appropriate to the "
             f"role, understated. Flat-ish editorial vector-illustration "
             f"feel, soft grain, no photo-realism."),
            ("Negative prompt", NEG),
            ("Output spec", "Square 1:1, ≥ 800×800, transparent or "
             "neutral-paper background; consistent lighting/style across all "
             "four portraits in the mission so the stakeholder grid feels one set."),
        ],
        "qa": [
            "Reads as a composite role, not an identifiable individual.",
            "Dignity check: competent and human, no poverty-porn.",
            "Style consistent with the other 3 portraits in the same mission.",
            "Placed at the exact path (or extension-only change made in the "
            "content module).",
        ],
    }


def image_block(code, title, tool, path, prompt, spec):
    return {
        "title": f"{code}. {title}",
        "rows": [
            ("Tool", tool),
            ("WIRE TO (exact path)", path),
            ("Prompt", prompt),
            ("Negative prompt", NEG),
            ("Output spec", spec),
        ],
        "qa": [
            "No fabricated precise statistics rendered as if real — numbers "
            "shown must be the pedagogical/illustrative values from the "
            "production master, labelled as such where visible.",
            "Legible at the size it renders in the dossier (test in-page).",
            "Placed at the exact path; filename unchanged.",
        ],
    }


# ── Assemble the document model ──
def build_model():
    return {
        "title": "FUTUREPROOF — Asset Production Prompt Master (v2)",
        "subtitle": "The two fully-built missions: Khon Kaen (SDG 6) & "
                    "Chiang Mai (SDG 13). Every prompt maps 1:1 to the exact "
                    "path the code already expects.",
        "date": str(date.today()),
        "howto": [
            "WIRING CONTRACT: filenames and paths in this document are "
            "load-bearing. Produce the asset, drop it at the stated path, and "
            "it appears in the mission with NO code change. Do not rename.",
            "AUDIO SCRIPTS ARE VERBATIM: every spoken word must match the "
            "script exactly — it is the same text shown on screen and in the "
            ".vtt caption. Paraphrasing breaks the caption sync and the "
            "comprehension items that quote stakeholders.",
            "STATUS — Khon Kaen: code wired; .vtt captions already shipped; "
            "the 4 stakeholder MP3s are MISSING (placeholder UI shows until "
            "they exist); portraits are baseline SVGs you may upgrade.",
            "STATUS — Chiang Mai: code wired; NOTHING on disk yet — it needs "
            "the 4 MP3s, the 4 matching .vtt captions, and 4 portraits. The "
            ".vtt cue text is given here so captions match the audio exactly.",
            "GROUNDING DISCIPLINE: composite roles, never real individuals; "
            "no invented precise statistics; Thai context with dignity. This "
            "mirrors the scenario notice and the DPIA.",
            "Audio engine behaviour: the player auto-detects a missing file "
            "and shows a graceful 'production pending' chip — so partial "
            "delivery never breaks a demo. Ship audio as you finish each one.",
        ],
        "missions": [
            {
                "name": "MISSION 1 — “The Aquifer Below Khon Kaen” (SDG 6, Northeast)",
                "intro": "Status: wired, captions shipped, MP3 audio MISSING. "
                         "Produce A1–A4 first (highest visible impact).",
                "audio": [audio_block(s, "sdg06-khonkaen") for s in KK],
                "extra_audio_note": (
                    "Optional secondary audio already has shipped .vtt scripts "
                    "in assets/scenarios/sdg06-khonkaen/audio/ "
                    "(crisis-dispatch.vtt, tribunal-open.vtt) and full "
                    "voice-direction in scenarios/sdg06-khonkaen-aquifer-v1/"
                    "PRODUCTION-PROMPTS.md §A5–A6 — reuse those scripts "
                    "verbatim; not repeated here to avoid drift."),
                "portraits": [
                    portrait_block(1, "Smallholder rice farmer",
                        "/assets/scenarios/sdg06-khonkaen/images/stakeholder-01.svg",
                        "vulnerable", "warm ochre"),
                    portrait_block(2, "PWA Operations Manager",
                        "/assets/scenarios/sdg06-khonkaen/images/stakeholder-02.svg",
                        "institutional", "soft steel-blue"),
                    portrait_block(3, "Food-processing Production Director",
                        "/assets/scenarios/sdg06-khonkaen/images/stakeholder-03.svg",
                        "private", "warm bronze"),
                    portrait_block(4, "Community Health Volunteer (อสม.)",
                        "/assets/scenarios/sdg06-khonkaen/images/stakeholder-04.svg",
                        "vulnerable-adjacent", "sage green"),
                ],
                "images": [
                    image_block("C1", "Hero image (journey card / briefing)",
                        "ChatGPT Image / Midjourney",
                        "/assets/scenarios/sdg06-khonkaen/images/hero.webp",
                        "Editorial wide image: the Khorat Plateau dry-season "
                        "landscape above Khon Kaen — a cracked paddy edge and a "
                        "deep municipal wellhead in the mid-distance under a "
                        "high hot sky; restrained, cinematic, console-editorial "
                        "palette (obsidian/bronze/bone), no people foreground, "
                        "no text.",
                        "16:9, ≥ 1600px wide, .webp; muted so gold UI text "
                        "stays legible over it. (Optional — SVG baseline "
                        "already ships; this is an upgrade.)"),
                    image_block("C2", "Aquifer cross-section (dossier diagram)",
                        "Illustrator / SVG (keep vector)",
                        "/assets/scenarios/sdg06-khonkaen/images/aquifer-crosssection.svg",
                        "Clean labelled cross-section: surface, upper "
                        "recharge layer with shallow household wells, deep "
                        "confined layer with the municipal wellfield, a "
                        "rock-salt zone below; arrows showing drawdown and "
                        "potential saline intrusion. Editorial, restrained, "
                        "bilingual-ready labels. Numbers illustrative only.",
                        "SVG, viewBox ~ 1000×640, legible at 720px. Baseline "
                        "already ships — upgrade for clarity only."),
                    image_block("C3", "Drawdown chart (data-interpretation)",
                        "Vega/Observable → export SVG",
                        "/assets/scenarios/sdg06-khonkaen/images/drawdown-chart.svg",
                        "Simple line/area chart: water-table depth over ~10 "
                        "years trending down, three drought years marked. Axis "
                        "labelled; a visible footnote 'Illustrative pedagogical "
                        "figures — not verified field data.'",
                        "SVG, legible at 680px; colour-blind-safe."),
                ],
            },
            {
                "name": "MISSION 2 — “The Burning Season” (SDG 13, North / Chiang Mai)",
                "intro": "Status: wired, NOTHING on disk. Needs 4 MP3 + 4 .vtt "
                         "+ 4 portraits. The .vtt cue text is generated by this "
                         "build at the exact path (see scripts) so you only "
                         "produce audio that matches it.",
                "audio": [audio_block(s, "sdg13-chiangmai") for s in CM],
                "extra_audio_note": (
                    "No secondary audio for v1 of this mission. If a crisis "
                    "dispatch is added later, follow the Khon Kaen A5 pattern."),
                "portraits": [
                    portrait_block(1, "Upland maize smallholder",
                        "/assets/scenarios/sdg13-chiangmai/images/stakeholder-01.svg",
                        "vulnerable", "warm ochre / earth"),
                    portrait_block(2, "Respiratory clinician",
                        "/assets/scenarios/sdg13-chiangmai/images/stakeholder-02.svg",
                        "institutional", "clean steel-blue / white"),
                    portrait_block(3, "Agribusiness procurement manager",
                        "/assets/scenarios/sdg13-chiangmai/images/stakeholder-03.svg",
                        "private", "warm bronze / corporate slate"),
                    portrait_block(4, "District enforcement officer",
                        "/assets/scenarios/sdg13-chiangmai/images/stakeholder-04.svg",
                        "institutional", "muted sage / field khaki"),
                ],
                "images": [
                    image_block("C1", "Hero image (journey card / briefing)",
                        "ChatGPT Image / Midjourney",
                        "/assets/scenarios/sdg13-chiangmai/images/hero.webp",
                        "Editorial wide image: the Ping valley under grey "
                        "burning-season haze, Doi Suthep a faint silhouette, "
                        "an upland field edge with thin smoke at the treeline; "
                        "console-editorial palette, no people foreground, no "
                        "text, restrained and cinematic.",
                        "16:9, ≥ 1600px, .webp; muted for legible gold UI text."),
                    image_block("C2", "Temperature-inversion diagram (dossier)",
                        "Illustrator / SVG",
                        "/assets/scenarios/sdg13-chiangmai/images/inversion-diagram.svg",
                        "Labelled cross-section of the Ping basin: cool "
                        "polluted air pooled in the valley, a warm 'lid' layer "
                        "above trapping it, smoke sources (field + "
                        "transboundary arrow at the border). Editorial, "
                        "bilingual-ready labels, illustrative only.",
                        "SVG, viewBox ~ 1000×640, legible at 720px."),
                    image_block("C3", "PM2.5 season chart (data-interpretation)",
                        "Vega/Observable → export SVG",
                        "/assets/scenarios/sdg13-chiangmai/images/pm25-chart.svg",
                        "Line chart: PM2.5 across a year, sharp Feb–Apr spike "
                        "into hazardous band; a marked WHO-guideline reference "
                        "line; footnote 'Illustrative pedagogical figures — "
                        "shape is realistic, exact values are not field data.'",
                        "SVG, legible at 680px; colour-blind-safe."),
                ],
            },
        ],
        # Chiang Mai .vtt cue text (this is what scripts/wire-chiangmai-captions
        # writes to disk so captions match the produced audio exactly).
        "captions": CM,
    }


# ── Renderers ───────────────────────────────────────────────────────────
def render_md(m):
    L = []
    L.append(f"# {m['title']}\n")
    L.append(f"*{m['subtitle']}*  ·  generated {m['date']}\n")
    L.append("> Single source of truth: `scripts/gen-production-prompts.py`. "
             "Regenerate with `python3 scripts/gen-production-prompts.py`.\n")
    L.append("## How to use this document\n")
    for h in m["howto"]:
        L.append(f"- {h}")
    L.append("")
    for ms in m["missions"]:
        L.append(f"\n---\n\n# {ms['name']}\n")
        L.append(f"_{ms['intro']}_\n")
        L.append("## Part A — Audio (ElevenLabs)\n")
        for a in ms["audio"]:
            L.append(f"### {a['title']}\n")
            for k, v in a["rows"]:
                L.append(f"- **{k}:** {v}")
            L.append("\n**SCRIPT (verbatim — speak exactly this):**\n")
            L.append(f"> {a['script']}\n")
            L.append("**QA checklist:**")
            for q in a["qa"]:
                L.append(f"- [ ] {q}")
            L.append("")
        L.append(f"_{ms['extra_audio_note']}_\n")
        L.append("## Part B — Stakeholder portraits\n")
        for b in ms["portraits"]:
            L.append(f"### {b['title']}\n")
            for k, v in b["rows"]:
                L.append(f"- **{k}:** {v}")
            L.append("\n**QA:**")
            for q in b["qa"]:
                L.append(f"- [ ] {q}")
            L.append("")
        L.append("## Part C — Scene & data images\n")
        for c in ms["images"]:
            L.append(f"### {c['title']}\n")
            for k, v in c["rows"]:
                L.append(f"- **{k}:** {v}")
            L.append("\n**QA:**")
            for q in c["qa"]:
                L.append(f"- [ ] {q}")
            L.append("")
    L.append("\n---\n\n## Appendix — Chiang Mai caption (.vtt) cue text\n")
    L.append("These are written to disk by `scripts/gen-production-prompts.py` "
             "at the exact `.vtt` paths so the produced audio only has to "
             "match this text:\n")
    for s in m["captions"]:
        L.append(f"**{s['vtt']}**\n")
        L.append(f"> {s['script']}\n")
    return "\n".join(L)


def render_docx(m):
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    base = doc.styles["Normal"]
    base.font.name = "Calibri"
    base.font.size = Pt(10.5)

    BRZ = RGBColor(0xB5, 0x7C, 0x34)
    INK = RGBColor(0x07, 0x17, 0x34)

    t = doc.add_paragraph()
    r = t.add_run(m["title"])
    r.bold = True
    r.font.size = Pt(22)
    r.font.color.rgb = INK
    s = doc.add_paragraph()
    rs = s.add_run(m["subtitle"])
    rs.italic = True
    rs.font.size = Pt(11)
    d = doc.add_paragraph()
    d.add_run(f"Generated {m['date']} · regenerate with "
              f"scripts/gen-production-prompts.py").font.size = Pt(8.5)

    h = doc.add_paragraph()
    hr = h.add_run("How to use this document")
    hr.bold = True
    hr.font.size = Pt(14)
    hr.font.color.rgb = BRZ
    for item in m["howto"]:
        doc.add_paragraph(item, style="List Bullet")

    def section_head(text, size=16):
        doc.add_page_break()
        p = doc.add_paragraph()
        rr = p.add_run(text)
        rr.bold = True
        rr.font.size = Pt(size)
        rr.font.color.rgb = INK

    def sub_head(text):
        p = doc.add_paragraph()
        rr = p.add_run(text)
        rr.bold = True
        rr.font.size = Pt(12)
        rr.font.color.rgb = BRZ

    def asset(a, script_label=None):
        p = doc.add_paragraph()
        rr = p.add_run(a["title"])
        rr.bold = True
        rr.font.size = Pt(12.5)
        tbl = doc.add_table(rows=0, cols=2)
        tbl.style = "Light Grid Accent 1"
        for k, v in a["rows"]:
            row = tbl.add_row().cells
            kr = row[0].paragraphs[0].add_run(k)
            kr.bold = True
            kr.font.size = Pt(9)
            vp = row[1].paragraphs[0].add_run(str(v))
            vp.font.size = Pt(9)
        if "script" in a:
            sp = doc.add_paragraph()
            spr = sp.add_run(script_label or "SCRIPT — speak exactly this:")
            spr.bold = True
            spr.font.size = Pt(9.5)
            spr.font.color.rgb = BRZ
            q = doc.add_paragraph()
            q.paragraph_format.left_indent = Inches(0.3)
            qr = q.add_run(a["script"])
            qr.italic = True
            qr.font.size = Pt(11)
        if "qa" in a:
            qh = doc.add_paragraph()
            qhr = qh.add_run("QA checklist")
            qhr.bold = True
            qhr.font.size = Pt(9)
            for qq in a["qa"]:
                doc.add_paragraph(qq, style="List Bullet 2"
                                  if "List Bullet 2" in doc.styles else "List Bullet")
        doc.add_paragraph("")

    for ms in m["missions"]:
        section_head(ms["name"])
        intro = doc.add_paragraph()
        ir = intro.add_run(ms["intro"])
        ir.italic = True
        ir.font.size = Pt(10)
        sub_head("Part A — Audio (ElevenLabs)")
        for a in ms["audio"]:
            asset(a)
        nn = doc.add_paragraph()
        nr = nn.add_run(ms["extra_audio_note"])
        nr.italic = True
        nr.font.size = Pt(9)
        sub_head("Part B — Stakeholder portraits")
        for b in ms["portraits"]:
            asset(b)
        sub_head("Part C — Scene & data images")
        for c in ms["images"]:
            asset(c)

    section_head("Appendix — Chiang Mai caption (.vtt) cue text", 14)
    doc.add_paragraph(
        "Written to disk at the exact .vtt paths by the build script so the "
        "produced audio only has to match this text verbatim.")
    for sx in m["captions"]:
        ph = doc.add_paragraph()
        phr = ph.add_run(sx["vtt"])
        phr.bold = True
        phr.font.size = Pt(9.5)
        pq = doc.add_paragraph()
        pq.paragraph_format.left_indent = Inches(0.3)
        pqr = pq.add_run(sx["script"])
        pqr.italic = True
        pqr.font.size = Pt(10.5)

    doc.save(OUT_DOCX)


def write_chiangmai_vtt(m):
    """Write WebVTT caption files so produced audio drops in matched."""
    for s in m["captions"]:
        rel = s["vtt"].lstrip("/")
        dest = os.path.join(ROOT, rel)
        os.makedirs(os.path.dirname(dest), exist_ok=True)
        end = "00:%02d:%02d.000" % (s["dur"] // 60, s["dur"] % 60)
        body = ("WEBVTT\n\nNOTE Auto-generated from the wired transcript "
                "(scripts/gen-production-prompts.py). The produced MP3 must "
                "match this text verbatim.\n\n"
                "1\n00:00:00.000 --> %s\n%s\n" % (end, s["script"]))
        with open(dest, "w", encoding="utf-8") as fh:
            fh.write(body)


def main():
    model = build_model()
    os.makedirs(os.path.dirname(OUT_MD), exist_ok=True)
    with open(OUT_MD, "w", encoding="utf-8") as fh:
        fh.write(render_md(model))
    print("wrote", os.path.relpath(OUT_MD, ROOT))
    write_chiangmai_vtt(model)
    print("wrote Chiang Mai .vtt captions (4)")
    try:
        render_docx(model)
        print("wrote", os.path.relpath(OUT_DOCX, ROOT))
    except ImportError:
        print("python-docx not installed — .md + .vtt written; "
              "run: pip install python-docx, then re-run for the .docx")


if __name__ == "__main__":
    main()
