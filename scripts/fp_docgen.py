#!/usr/bin/env python3
"""
fp_docgen.py — FUTUREPROOF editorial Word-document template + helpers.

Design intent: documents that read as hand-authored in Microsoft Word, in a
professional editorial template that echoes the FUTUREPROOF portal.

  • Headings  : Baskerville (macOS editorial serif ≈ portal's Fraunces/Cormorant)
  • Body      : Avenir Next (humanist sans, big & readable, near-black)
  • Labels    : Menlo (letter-spaced caps ≈ portal's JetBrains-Mono console tags)
  • Thai      : Sukhumvit Set (clean macOS Thai, for the bilingual manual)
  • Colours   : body = black #000000 · titles = Dark Red #7B1B1B · section heads
                & labels = Dark Gold #9C6F1C · hairlines = #C9A961

All public helpers build native Word constructs (styles, shading, borders,
PAGE fields) so the exported PDF looks authored, not generated.
"""

from docx import Document
from docx.shared import Pt, Mm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os, subprocess, tempfile, time

# ── palette / fonts ──────────────────────────────────────────────
SERIF, SANS, MONO, THAI = "Baskerville", "Avenir Next", "Menlo", "Sukhumvit Set"
INK   = RGBColor(0x00, 0x00, 0x00)   # body — black
DRED  = RGBColor(0x7B, 0x1B, 0x1B)   # dark red — titles
DGOLD = RGBColor(0x9C, 0x6F, 0x1C)   # dark gold — section heads / labels
GRAY  = RGBColor(0x5A, 0x5A, 0x60)   # secondary
GOLDH = "C9A961"                     # hairline gold (fill/borders, no '#')
REDH  = "7B1B1B"
GOLDFAINT = "F3ECDD"                  # callout gold wash
REDFAINT  = "F3E4E2"                  # callout red wash
HEADFILL  = "F4F1EA"                  # table header fill

# ── low-level XML helpers ────────────────────────────────────────
def _rpr(run):
    return run._element.get_or_add_rPr()

def set_fonts(run, ascii_=SANS, cs=None, ea=None):
    """Set ascii/hAnsi + complex-script + east-asian font for one run."""
    run.font.name = ascii_
    rpr = _rpr(run)
    rfonts = rpr.find(qn('w:rFonts'))
    if rfonts is None:
        rfonts = OxmlElement('w:rFonts'); rpr.append(rfonts)
    rfonts.set(qn('w:ascii'), ascii_)
    rfonts.set(qn('w:hAnsi'), ascii_)
    rfonts.set(qn('w:cs'), cs or ascii_)
    if ea:
        rfonts.set(qn('w:eastAsia'), ea)

def letter_space(run, pts):
    rpr = _rpr(run)
    sp = OxmlElement('w:spacing'); sp.set(qn('w:val'), str(int(pts * 20))); rpr.append(sp)

def _shade_elem(parent_pr, hexfill):
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), hexfill)
    parent_pr.append(shd)

def shade_paragraph(p, hexfill):
    _shade_elem(p._p.get_or_add_pPr(), hexfill)

def shade_cell(cell, hexfill):
    _shade_elem(cell._tc.get_or_add_tcPr(), hexfill)

def left_border(p, hexcol, sz=18):
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement('w:pBdr')
    for edge in ('left',):
        b = OxmlElement(f'w:{edge}')
        b.set(qn('w:val'), 'single'); b.set(qn('w:sz'), str(sz))
        b.set(qn('w:space'), '8'); b.set(qn('w:color'), hexcol)
        pbdr.append(b)
    pPr.append(pbdr)

def bottom_rule(p, hexcol, sz=6):
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement('w:pBdr')
    b = OxmlElement('w:bottom')
    b.set(qn('w:val'), 'single'); b.set(qn('w:sz'), str(sz))
    b.set(qn('w:space'), '4'); b.set(qn('w:color'), hexcol)
    pbdr.append(b); pPr.append(pbdr)

def _add_run(p, text, *, font=SANS, size=12.5, color=INK, bold=False,
             italic=False, ls=None, cs=None, ea=None):
    r = p.add_run(text)
    set_fonts(r, font, cs=cs, ea=ea)
    r.font.size = Pt(size); r.font.color.rgb = color
    r.font.bold = bold; r.font.italic = italic
    if ls: letter_space(r, ls)
    return r

def _markup_runs(p, text, *, font=SANS, size=12.5, color=INK):
    """Render a string with **bold** segments into runs."""
    parts = text.split('**')
    for i, seg in enumerate(parts):
        if seg == '':
            continue
        _add_run(p, seg, font=font, size=size, color=color, bold=(i % 2 == 1))

# ── document factory + styles ────────────────────────────────────
def new_doc():
    doc = Document()
    sec = doc.sections[0]
    sec.page_width, sec.page_height = Mm(210), Mm(297)
    sec.top_margin = Mm(22); sec.bottom_margin = Mm(20)
    sec.left_margin = Mm(22); sec.right_margin = Mm(22)
    normal = doc.styles['Normal']
    normal.font.name = SANS; normal.font.size = Pt(12.5); normal.font.color.rgb = INK
    rpr = normal.element.get_or_add_rPr()
    rfonts = rpr.find(qn('w:rFonts'))
    if rfonts is None:
        rfonts = OxmlElement('w:rFonts'); rpr.append(rfonts)
    rfonts.set(qn('w:ascii'), SANS); rfonts.set(qn('w:hAnsi'), SANS)
    rfonts.set(qn('w:cs'), SANS); rfonts.set(qn('w:eastAsia'), THAI)
    return doc

# ── public content helpers ───────────────────────────────────────
def kicker(doc, text):
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(2)
    # cs/ea = Thai so any Thai in the label renders in Sukhumvit, Latin stays Menlo
    _add_run(p, text.upper(), font=MONO, size=9, color=DGOLD, ls=1.6, cs=THAI, ea=THAI)
    return p

def title(doc, text):
    p = doc.add_paragraph(); pf = p.paragraph_format
    pf.space_before = Pt(2); pf.space_after = Pt(4); pf.line_spacing = 1.02
    _add_run(p, text, font=SERIF, size=27, color=DRED, bold=False)
    return p

def thai_sub(doc, text):
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(10)
    _add_run(p, text, font=THAI, size=13, color=DGOLD, cs=THAI, ea=THAI)
    return p

def lede(doc, text):
    p = doc.add_paragraph(); pf = p.paragraph_format
    pf.space_after = Pt(14); pf.line_spacing = 1.3
    _add_run(p, text, font=SERIF, size=14.5, color=GRAY, italic=True)
    return p

def h2(doc, text):
    p = doc.add_paragraph(); pf = p.paragraph_format
    pf.space_before = Pt(16); pf.space_after = Pt(7)
    _add_run(p, text, font=SERIF, size=16.5, color=DGOLD, bold=False)
    bottom_rule(p, GOLDH, sz=6)
    return p

def h3(doc, text):
    p = doc.add_paragraph(); pf = p.paragraph_format
    pf.space_before = Pt(10); pf.space_after = Pt(3)
    _add_run(p, text, font=SERIF, size=13.5, color=INK, bold=True)
    return p

def body(doc, text):
    p = doc.add_paragraph(); pf = p.paragraph_format
    pf.space_after = Pt(7); pf.line_spacing = 1.35
    _markup_runs(p, text, font=SANS, size=12.5, color=INK)
    return p

def thai_body(doc, text):
    p = doc.add_paragraph(); pf = p.paragraph_format
    pf.space_after = Pt(7); pf.line_spacing = 1.45
    _add_run(p, text, font=THAI, size=12.5, color=INK, cs=THAI, ea=THAI)
    return p

def bullets(doc, items):
    for it in items:
        p = doc.add_paragraph(style='List Bullet'); pf = p.paragraph_format
        pf.space_after = Pt(3); pf.line_spacing = 1.3
        _markup_runs(p, it, font=SANS, size=12.5, color=INK)

def callout(doc, label, text, kind='gold'):
    fill = REDFAINT if kind == 'red' else GOLDFAINT
    bcol = REDH if kind == 'red' else GOLDH
    lcol = DRED if kind == 'red' else DGOLD
    pl = doc.add_paragraph(); pl.paragraph_format.space_before = Pt(8)
    pl.paragraph_format.space_after = Pt(0)
    shade_paragraph(pl, fill); left_border(pl, bcol, sz=22)
    _add_run(pl, label.upper(), font=MONO, size=8.5, color=lcol, ls=1.4, bold=True)
    pb = doc.add_paragraph(); pb.paragraph_format.space_after = Pt(8)
    pb.paragraph_format.line_spacing = 1.3
    shade_paragraph(pb, fill); left_border(pb, bcol, sz=22)
    _markup_runs(pb, text, font=SANS, size=11.5, color=INK)

def divider(doc):
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4); bottom_rule(p, GOLDH, sz=4)

def _style_table_borders(tbl, hexcol="D9D2C4", sz=4):
    tblPr = tbl._tbl.tblPr
    borders = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        e = OxmlElement(f'w:{edge}')
        e.set(qn('w:val'), 'single'); e.set(qn('w:sz'), str(sz))
        e.set(qn('w:space'), '0'); e.set(qn('w:color'), hexcol)
        borders.append(e)
    tblPr.append(borders)

def kv_table(doc, rows):
    """2-column key/value table (mono gold key, sans value)."""
    tbl = doc.add_table(rows=0, cols=2); tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    _style_table_borders(tbl)
    tbl.columns[0].width = Mm(52); tbl.columns[1].width = Mm(114)
    for k, v in rows:
        row = tbl.add_row().cells
        row[0].width = Mm(52); row[1].width = Mm(114)
        pk = row[0].paragraphs[0]; pk.paragraph_format.space_after = Pt(2)
        _add_run(pk, k.upper(), font=MONO, size=8.5, color=DGOLD, ls=0.6)
        pv = row[1].paragraphs[0]; pv.paragraph_format.space_after = Pt(2)
        pv.paragraph_format.line_spacing = 1.3
        _markup_runs(pv, v, font=SANS, size=11.5, color=INK)
    return tbl

def data_table(doc, headers, rows, widths_mm=None):
    """Header-row table: mono gold headers on faint fill, sans body."""
    ncol = len(headers)
    tbl = doc.add_table(rows=1, cols=ncol); tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    _style_table_borders(tbl)
    hdr = tbl.rows[0].cells
    for i, h in enumerate(headers):
        shade_cell(hdr[i], HEADFILL)
        p = hdr[i].paragraphs[0]; p.paragraph_format.space_after = Pt(1)
        _add_run(p, str(h).upper(), font=MONO, size=8, color=DGOLD, ls=0.8, bold=True)
    for r in rows:
        cells = tbl.add_row().cells
        for i, val in enumerate(r):
            p = cells[i].paragraphs[0]; p.paragraph_format.space_after = Pt(1)
            p.paragraph_format.line_spacing = 1.25
            _markup_runs(p, str(val), font=SANS, size=10.5, color=INK)
    if widths_mm:
        for i, w in enumerate(widths_mm):
            for row in tbl.rows:
                row.cells[i].width = Mm(w)
    return tbl

def footer(doc, *_ignored):
    """Standard FUTUREPROOF footer: authorship line + concise CC BY-NC notice.
    (Extra positional args are accepted and ignored for backward compatibility.)"""
    sec = doc.sections[0]
    f = sec.footer; f.is_linked_to_previous = False
    # line 1 — authorship
    p1 = f.paragraphs[0]; p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p1.paragraph_format.space_after = Pt(1)
    _add_run(p1, "Dr. Payungsak Kaenchan  |  Faculty of Liberal Arts, Mahidol University",
             font=MONO, size=7.5, color=GRAY, ls=0.5)
    # line 2 — license / reuse notice
    p2 = f.add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_run(p2, "Creative Commons BY-NC · Educational, non-commercial use only · "
                 "No profit · Reuse or duplication must credit the author and obtain prior consent.",
             font=MONO, size=6.8, color=GRAY, ls=0.3)

def cover(doc, kicker_text, title_text, subtitle, thai_line, slug, image_dir=None):
    """Cover page: top banner image (if present) + title block. New page after."""
    img = None
    if image_dir:
        for ext in ('.png', '.jpg', '.jpeg', '.webp'):
            cand = os.path.join(image_dir, slug + ext)
            if os.path.exists(cand):
                img = cand; break
    if img:
        pic = doc.add_paragraph(); pic.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pic.paragraph_format.space_after = Pt(18)
        run = pic.add_run(); run.add_picture(img, width=Mm(166))
    else:
        ph = doc.add_paragraph(); ph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        ph.paragraph_format.space_before = Pt(40); ph.paragraph_format.space_after = Pt(40)
        shade_paragraph(ph, GOLDFAINT)
        _add_run(ph, "[ COVER IMAGE — drop %s.png in /docs/document-design/covers/ ]" % slug,
                 font=MONO, size=9, color=DGOLD, ls=0.4)
    k = doc.add_paragraph(); k.alignment = WD_ALIGN_PARAGRAPH.CENTER
    k.paragraph_format.space_after = Pt(4)
    _add_run(k, kicker_text.upper(), font=MONO, size=10, color=DGOLD, ls=1.8, cs=THAI, ea=THAI)
    t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t.paragraph_format.space_after = Pt(6); t.paragraph_format.line_spacing = 1.02
    _add_run(t, title_text, font=SERIF, size=32, color=DRED)
    if subtitle:
        s = doc.add_paragraph(); s.alignment = WD_ALIGN_PARAGRAPH.CENTER
        s.paragraph_format.space_after = Pt(4)
        _add_run(s, subtitle, font=SERIF, size=15, color=GRAY, italic=True)
    if thai_line:
        th = doc.add_paragraph(); th.alignment = WD_ALIGN_PARAGRAPH.CENTER
        th.paragraph_format.space_after = Pt(24)
        _add_run(th, thai_line, font=THAI, size=13, color=DGOLD, cs=THAI, ea=THAI)
    auth = doc.add_paragraph(); auth.alignment = WD_ALIGN_PARAGRAPH.CENTER
    auth.paragraph_format.space_before = Pt(20)
    _add_run(auth, "Dr. Payungsak Kaenchan  ·  Faculty of Liberal Arts, Mahidol University",
             font=SANS, size=11, color=INK)
    sub = doc.add_paragraph(); sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_run(sub, "Submitted to the SPU Tech Creative Learning Awards 2569",
             font=MONO, size=8, color=GRAY, ls=0.8)
    doc.add_page_break()

# ── docx → PDF via Microsoft Word (AppleScript) ──────────────────
def to_pdf(docx_path, pdf_path=None):
    docx_path = os.path.abspath(docx_path)
    pdf_path = os.path.abspath(pdf_path or os.path.splitext(docx_path)[0] + ".pdf")
    # Isolate each conversion: a stale/open Word session causes AppleEvent
    # timeouts, so start from a clean state every time.
    subprocess.run(["pkill", "-x", "Microsoft Word"], capture_output=True)
    time.sleep(2)
    script = f'''
    tell application "Microsoft Word"
        activate
        open POSIX file "{docx_path}"
        delay 2
        tell active document
            with timeout of 600 seconds
                save as it file name "{pdf_path}" file format format PDF
            end timeout
        end tell
        close active document saving no
    end tell
    '''
    res = subprocess.run(["osascript", "-e", script], capture_output=True, text=True)
    if res.returncode != 0:
        raise RuntimeError("Word PDF export failed:\n" + res.stderr)
    return pdf_path
